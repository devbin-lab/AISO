"""문서 파일 텍스트 추출 — 에이전트 read_file이 PDF/엑셀/워드/한글 등을 읽게 한다.

각 추출기는 텍스트를 돌려주거나 ExtractError를 던진다. (tools.py가 ToolError로 변환)
코드·md·csv·txt 등 순수 텍스트는 tools.read_file의 기본 경로가 처리한다.
"""

from __future__ import annotations

import re
import zlib
import zipfile
from dataclasses import dataclass
from html import escape
from pathlib import Path
from xml.etree import ElementTree as ET


class ExtractError(Exception):
    """문서 추출 실패 — 사유는 모델에게 그대로 전달된다."""


@dataclass(frozen=True)
class ExtractSegment:
    """A small, addressable piece of a document.

    ``text`` remains the original extracted wording.  The evidence-ToDo flow
    stores this wording verbatim rather than asking a model to recreate it,
    which makes every generated task traceable to a page, slide, paragraph or
    spreadsheet row.
    """

    location: str
    text: str


# ---------- PDF ----------

def extract_pdf(target: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractError("PDF 읽기 라이브러리(pypdf)가 설치되어 있지 않습니다.")
    try:
        reader = PdfReader(str(target))
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"PDF를 열 수 없습니다: {e}")
    total = len(reader.pages)
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        if i >= 60:
            parts.append(f"\n…(총 {total}쪽 중 60쪽까지만 표시)")
            break
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\n\n".join(parts)


def extract_pdf_segments(target: Path) -> list[ExtractSegment]:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractError("PDF 읽기 라이브러리(pypdf)가 설치되어 있지 않습니다.")
    try:
        reader = PdfReader(str(target))
    except Exception as error:  # noqa: BLE001
        raise ExtractError(f"PDF를 열 수 없습니다: {error}") from error
    segments: list[ExtractSegment] = []
    for index, page in enumerate(reader.pages[:60], 1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001
            text = ""
        if text:
            segments.append(ExtractSegment(location=f"{index}쪽", text=text))
    return segments


# ---------- 엑셀 (XLSX/XLSM) ----------

def extract_xlsx(target: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ExtractError("엑셀 읽기 라이브러리(openpyxl)가 설치되어 있지 않습니다.")
    try:
        wb = load_workbook(str(target), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"엑셀을 열 수 없습니다: {e}")
    out: list[str] = []
    try:
        for ws in wb.worksheets:
            out.append(f"# 시트: {ws.title}")
            for r, row in enumerate(ws.iter_rows(values_only=True)):
                if r >= 500:
                    out.append("…(행 500개 초과 생략)")
                    break
                cells = [("" if c is None else str(c)) for c in row]
                if any(c.strip() for c in cells):
                    out.append("\t".join(cells))
    finally:
        wb.close()
    return "\n".join(out)


def extract_xlsx_segments(target: Path) -> list[ExtractSegment]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ExtractError("엑셀 읽기 라이브러리(openpyxl)가 설치되어 있지 않습니다.")
    try:
        workbook = load_workbook(str(target), read_only=True, data_only=True)
    except Exception as error:  # noqa: BLE001
        raise ExtractError(f"엑셀을 열 수 없습니다: {error}") from error
    segments: list[ExtractSegment] = []
    try:
        for worksheet in workbook.worksheets:
            for row_number, row in enumerate(worksheet.iter_rows(values_only=True), 1):
                if row_number > 500:
                    break
                values = [("" if value is None else str(value)).strip() for value in row]
                text = "\t".join(values).strip()
                if text:
                    segments.append(ExtractSegment(location=f"{worksheet.title} · {row_number}행", text=text))
    finally:
        workbook.close()
    return segments


def extract_xls(target: Path) -> str:
    raise ExtractError("구형 .xls 형식은 지원하지 않습니다. 엑셀에서 .xlsx로 저장 후 다시 시도하세요.")


# ---------- 워드 (DOCX) ----------

def extract_docx(target: Path) -> str:
    try:
        import docx
    except ImportError:
        raise ExtractError("워드 읽기 라이브러리(python-docx)가 설치되어 있지 않습니다.")
    try:
        d = docx.Document(str(target))
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"워드 문서를 열 수 없습니다: {e}")
    parts: list[str] = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_docx_segments(target: Path) -> list[ExtractSegment]:
    try:
        import docx
    except ImportError:
        raise ExtractError("워드 읽기 라이브러리(python-docx)가 설치되어 있지 않습니다.")
    try:
        document = docx.Document(str(target))
    except Exception as error:  # noqa: BLE001
        raise ExtractError(f"워드 문서를 열 수 없습니다: {error}") from error
    segments: list[ExtractSegment] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            segments.append(ExtractSegment(location=f"문단 {index}", text=text))
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = "\t".join(cell.text.strip() for cell in row.cells).strip()
            if text:
                segments.append(ExtractSegment(location=f"표 {table_index} · {row_index}행", text=text))
    return segments


# ---------- PowerPoint (PPTX/PPTM) ----------

_PPT_DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _ppt_slide_number(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name, re.IGNORECASE)
    return int(match.group(1)) if match else 0


def extract_pptx(target: Path) -> str:
    """Extract visible text from OOXML PowerPoint slides without LibreOffice.

    A PPTX is a ZIP package.  Reading the slide XML directly keeps the Agent's
    document-reading path available in the portable Python runtime and gives
    each source location a stable slide number for later citation.
    """
    try:
        with zipfile.ZipFile(target) as archive:
            slide_names = sorted(
                (
                    name for name in archive.namelist()
                    if re.fullmatch(r"ppt/slides/slide\d+\.xml", name, re.IGNORECASE)
                ),
                key=_ppt_slide_number,
            )
            if not slide_names:
                raise ExtractError("PowerPoint 슬라이드 XML을 찾을 수 없습니다.")

            slides: list[str] = []
            paragraph_tag = f"{{{_PPT_DRAWING_NS}}}p"
            text_tag = f"{{{_PPT_DRAWING_NS}}}t"
            for slide_name in slide_names:
                try:
                    root = ET.fromstring(archive.read(slide_name))
                except (ET.ParseError, KeyError) as error:
                    raise ExtractError(f"PowerPoint 슬라이드를 읽을 수 없습니다: {slide_name}") from error

                paragraphs: list[str] = []
                for paragraph in root.iter(paragraph_tag):
                    text = "".join((node.text or "") for node in paragraph.iter(text_tag)).strip()
                    if text:
                        paragraphs.append(text)
                number = _ppt_slide_number(slide_name)
                body = "\n".join(paragraphs) or "[텍스트가 없는 슬라이드]"
                slides.append(f"# 슬라이드 {number}\n{body}")
    except zipfile.BadZipFile as error:
        raise ExtractError("PPTX/PPTM 파일이 올바른 Office 문서 형식이 아닙니다.") from error
    except OSError as error:
        raise ExtractError(f"PowerPoint 파일을 열 수 없습니다: {error}") from error
    return "\n\n".join(slides)


def extract_pptx_segments(target: Path) -> list[ExtractSegment]:
    """Return every slide as a separately addressable evidence segment."""
    try:
        with zipfile.ZipFile(target) as archive:
            slide_names = sorted(
                (name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name, re.IGNORECASE)),
                key=_ppt_slide_number,
            )
            if not slide_names:
                raise ExtractError("PowerPoint 슬라이드 XML을 찾을 수 없습니다.")
            segments: list[ExtractSegment] = []
            paragraph_tag = f"{{{_PPT_DRAWING_NS}}}p"
            text_tag = f"{{{_PPT_DRAWING_NS}}}t"
            for slide_name in slide_names:
                try:
                    root = ET.fromstring(archive.read(slide_name))
                except (ET.ParseError, KeyError) as error:
                    raise ExtractError(f"PowerPoint 슬라이드를 읽을 수 없습니다: {slide_name}") from error
                paragraphs = []
                for paragraph in root.iter(paragraph_tag):
                    text = "".join((node.text or "") for node in paragraph.iter(text_tag)).strip()
                    if text:
                        paragraphs.append(text)
                if paragraphs:
                    segments.append(ExtractSegment(location=f"슬라이드 {_ppt_slide_number(slide_name)}", text="\n".join(paragraphs)))
            return segments
    except zipfile.BadZipFile as error:
        raise ExtractError("PPTX/PPTM 파일이 올바른 Office 문서 형식이 아닙니다.") from error
    except OSError as error:
        raise ExtractError(f"PowerPoint 파일을 열 수 없습니다: {error}") from error


def pptx_to_html(target: Path, destination: Path) -> None:
    """Create a safe, text-first HTML rendition for offline review.

    This is intentionally a faithful *reading* rendition rather than a layout
    clone.  It gives the Agent and the user a durable HTML fallback when a
    PowerPoint renderer is unavailable, while preserving slide boundaries.
    """
    extracted = extract_pptx(target)
    sections: list[str] = []
    current_title = ""
    current_lines: list[str] = []
    for line in extracted.splitlines():
        if line.startswith("# 슬라이드 "):
            if current_title:
                sections.append(
                    f"<section><h2>{escape(current_title)}</h2><pre>{escape(chr(10).join(current_lines))}</pre></section>"
                )
            current_title = line[2:]
            current_lines = []
        else:
            current_lines.append(line)
    if current_title:
        sections.append(
            f"<section><h2>{escape(current_title)}</h2><pre>{escape(chr(10).join(current_lines))}</pre></section>"
        )
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(
        "<!doctype html><html lang=\"ko\"><meta charset=\"utf-8\">"
        f"<title>{escape(target.name)} - 텍스트 보기</title>"
        "<style>body{font-family:system-ui,sans-serif;max-width:960px;margin:40px auto;padding:0 20px;}"
        "section{border-bottom:1px solid #ddd;padding:0 0 24px;margin:0 0 24px;}"
        "pre{white-space:pre-wrap;font:inherit;line-height:1.6;}</style>"
        f"<h1>{escape(target.name)}</h1><p>PowerPoint 텍스트 추출 보기</p>{''.join(sections)}</html>",
        encoding="utf-8",
    )


# ---------- 한글 HWPX (ZIP + XML) ----------

def extract_hwpx(target: Path) -> str:
    import zipfile
    from xml.etree import ElementTree as ET

    parts: list[str] = []
    try:
        with zipfile.ZipFile(str(target)) as z:
            names = sorted(
                n for n in z.namelist() if n.startswith("Contents/") and n.endswith(".xml") and "section" in n.lower()
            )
            if not names:
                names = sorted(n for n in z.namelist() if n.endswith(".xml") and "section" in n.lower())
            for name in names:
                xml = z.read(name).decode("utf-8", errors="replace")
                try:
                    root = ET.fromstring(xml)
                    for el in root.iter():
                        tag = el.tag.rsplit("}", 1)[-1]
                        if tag == "p":
                            parts.append("\n")
                        elif tag == "t" and el.text:
                            parts.append(el.text)
                except ET.ParseError:
                    parts.append(re.sub(r"<[^>]+>", "", xml))
    except zipfile.BadZipFile:
        raise ExtractError("HWPX(ZIP) 파일을 열 수 없습니다.")
    return "".join(parts)


# ---------- 한글 HWP 5.0 (OLE 복합 파일) ----------

_EXT_CTRL = frozenset({1, 2, 3, 11, 12, 14, 15, 16, 17, 18, 21, 22, 23})
_INLINE_CTRL = frozenset({4, 5, 6, 7, 8, 9, 19, 20})
_HWPTAG_PARA_TEXT = 67


def _decode_hwp_text(data: bytes) -> str:
    out: list[str] = []
    i = 0
    n = len(data) - (len(data) % 2)
    while i < n:
        code = data[i] | (data[i + 1] << 8)
        if code < 32:
            if code in (10, 13):
                out.append("\n")
                i += 2
            elif code in _EXT_CTRL or code in _INLINE_CTRL:
                i += 16  # 확장/인라인 컨트롤은 8 wchar 차지
            else:
                i += 2
        else:
            out.append(chr(code))
            i += 2
    return "".join(out)


def _hwp_section_text(buf: bytes) -> str:
    out: list[str] = []
    p = 0
    n = len(buf)
    while p + 4 <= n:
        header = int.from_bytes(buf[p:p + 4], "little")
        p += 4
        tag = header & 0x3FF
        size = (header >> 20) & 0xFFF
        if size == 0xFFF:
            size = int.from_bytes(buf[p:p + 4], "little")
            p += 4
        data = buf[p:p + size]
        p += size
        if tag == _HWPTAG_PARA_TEXT:
            out.append(_decode_hwp_text(data))
    return "\n".join(out)


def extract_hwp(target: Path) -> str:
    try:
        import olefile
    except ImportError:
        raise ExtractError("한글(HWP) 읽기 라이브러리(olefile)가 설치되어 있지 않습니다.")
    if not olefile.isOleFile(str(target)):
        raise ExtractError("HWP 5.0 형식이 아닙니다 (구형 HWP 3.0 등은 미지원).")
    ole = olefile.OleFileIO(str(target))
    try:
        compressed = True
        if ole.exists("FileHeader"):
            fh = ole.openstream("FileHeader").read()
            if len(fh) > 36:
                compressed = bool(fh[36] & 0x01)
        sections = [
            e for e in ole.listdir()
            if len(e) == 2 and e[0] == "BodyText" and e[1].lower().startswith("section")
        ]
        sections.sort(key=lambda e: int((re.search(r"(\d+)", e[1]) or re.match("0", "0")).group()))
        out: list[str] = []
        for entry in sections:
            raw = ole.openstream(entry).read()
            if compressed:
                try:
                    raw = zlib.decompress(raw, -15)  # raw deflate
                except Exception:  # noqa: BLE001
                    continue
            out.append(_hwp_section_text(raw))
        text = "\n".join(out)
        if not text.strip():
            raise ExtractError("HWP에서 텍스트를 추출하지 못했습니다.")
        return text
    finally:
        ole.close()


# 확장자 → (추출기, 라벨)
EXTRACTORS = {
    ".pdf": (extract_pdf, "PDF"),
    ".pptx": (extract_pptx, "PowerPoint"),
    ".pptm": (extract_pptx, "PowerPoint"),
    ".xlsx": (extract_xlsx, "엑셀"),
    ".xlsm": (extract_xlsx, "엑셀"),
    ".xls": (extract_xls, "엑셀"),
    ".docx": (extract_docx, "워드"),
    ".hwp": (extract_hwp, "한글 HWP"),
    ".hwpx": (extract_hwpx, "한글 HWPX"),
}


def extract_document_segments(target: Path) -> list[ExtractSegment]:
    """Extract document text with a stable location for each source chunk.

    This complements ``EXTRACTORS`` instead of replacing it: the existing
    ``read_file`` tool keeps its compact string response, while creator tools
    get page/slide/paragraph-level evidence that can be shown to a user.
    """
    ext = target.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_segments(target)
    if ext in {".pptx", ".pptm"}:
        return extract_pptx_segments(target)
    if ext in {".docx"}:
        return extract_docx_segments(target)
    if ext in {".xlsx", ".xlsm"}:
        return extract_xlsx_segments(target)
    if ext in {".txt", ".md", ".csv"}:
        try:
            lines = target.read_text(encoding="utf-8", errors="replace").splitlines()
        except OSError as error:
            raise ExtractError(f"텍스트 파일을 열 수 없습니다: {error}") from error
        segments: list[ExtractSegment] = []
        for start in range(0, min(len(lines), 3000), 80):
            chunk = "\n".join(lines[start:start + 80]).strip()
            if chunk:
                end = min(start + 80, len(lines))
                segments.append(ExtractSegment(location=f"{start + 1}~{end}줄", text=chunk))
        return segments
    if ext in EXTRACTORS:
        extractor, _label = EXTRACTORS[ext]
        try:
            text = (extractor(target) or "").strip()
        except ExtractError:
            raise
        except Exception as error:  # noqa: BLE001
            raise ExtractError(f"문서 텍스트를 추출하지 못했습니다: {error}") from error
        return [ExtractSegment(location="본문", text=text)] if text else []
    raise ExtractError(f"지원하지 않는 문서 형식입니다: {target.suffix or '확장자 없음'}")
